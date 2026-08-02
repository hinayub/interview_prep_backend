"""Text extraction for uploaded resumes.

One public entry point, ``extract_text``. Callers never see pdfplumber or docx.
"""

import re
from pathlib import Path

import pdfplumber
from docx import Document


class UnsupportedFileType(Exception):
    """The upload is not a file type we can extract text from."""


class ParseError(Exception):
    """The file is a supported type but could not be read (corrupt/encrypted)."""


def _normalize(text: str) -> str:
    """Collapse the whitespace noise that PDF extraction always produces.

    Left in place, runs of blank lines and stray spaces waste a large share of the
    prompt budget on a 3B model with a small context window.
    """
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(file_obj) -> str:
    pages = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _extract_docx(file_obj) -> str:
    document = Document(file_obj)
    parts = [p.text for p in document.paragraphs]
    # Plenty of resumes lay out skills and dates in tables; skipping tables loses
    # exactly the structured content the matching agent cares most about.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


_EXTRACTORS = {".pdf": _extract_pdf, ".docx": _extract_docx}


def extract_text(uploaded_file) -> str:
    """Return normalized plain text from a Django ``UploadedFile``.

    Dispatches on the filename extension. Raises ``UnsupportedFileType`` for
    anything other than .pdf/.docx and ``ParseError`` if a supported file cannot
    be read.
    """
    suffix = Path(uploaded_file.name).suffix.lower()
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:
        supported = ", ".join(sorted(_EXTRACTORS))
        raise UnsupportedFileType(f"Cannot read '{suffix or uploaded_file.name}'. Supported: {supported}")

    uploaded_file.seek(0)
    try:
        text = extractor(uploaded_file)
    except Exception as exc:
        raise ParseError(f"Could not read this {suffix} file: {exc}") from exc
    finally:
        uploaded_file.seek(0)

    return _normalize(text)
