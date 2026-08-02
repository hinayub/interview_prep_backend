"""Shared test fixtures.

PDF and DOCX fixtures are generated at runtime rather than committed as binaries -
a real reportlab-produced PDF exercises the same pdfminer code path a user's resume
would, and there is nothing opaque in the repo to keep in sync.
"""

import io

import pytest
import requests
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from rest_framework.test import APIClient

from candidates.models import Candidate
from resumes.models import JobDescription, MatchAnalysis, Resume

# Long enough to clear RESUME_MIN_PARSED_CHARS (100) with room to spare.
SAMPLE_RESUME_LINES = [
    "Jane Q. Candidate",
    "Senior Backend Engineer - jane@example.com - +1 555 0100",
    "EXPERIENCE",
    "Acme Corp, Backend Engineer, 2020-2024.",
    "Built Django REST services handling 2M requests per day.",
    "Migrated a monolith to Celery-backed async task processing.",
    "SKILLS",
    "Python, Django, PostgreSQL, Redis, Docker, React, pytest.",
]

SAMPLE_JD_TEXT = (
    "We are hiring a Senior Python Engineer to own our Django REST API. "
    "You will design async pipelines with Celery, tune PostgreSQL queries, "
    "and mentor engineers. Requires 5+ years of Python and strong testing habits. "
    "Kubernetes and Terraform experience is a plus."
)


def build_pdf(lines=None, filename="resume.pdf") -> SimpleUploadedFile:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    y = 750
    for line in lines if lines is not None else SAMPLE_RESUME_LINES:
        pdf.drawString(72, y, line)
        y -= 18
    pdf.save()
    return SimpleUploadedFile(filename, buffer.getvalue(), content_type="application/pdf")


def build_docx(lines=None, filename="resume.docx") -> SimpleUploadedFile:
    buffer = io.BytesIO()
    document = Document()
    for line in lines if lines is not None else SAMPLE_RESUME_LINES:
        document.add_paragraph(line)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Certification"
    table.rows[0].cells[1].text = "AWS Solutions Architect"
    document.save(buffer)
    return SimpleUploadedFile(
        filename,
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@pytest.fixture
def pdf_resume():
    return build_pdf()


@pytest.fixture
def docx_resume():
    return build_docx()


@pytest.fixture
def make_candidate(db, django_user_model):
    def _make(username="jane", password="s3cret-passphrase"):
        user = django_user_model.objects.create_user(username=username, password=password)
        return Candidate.objects.create(user=user)

    return _make


@pytest.fixture
def candidate(make_candidate):
    return make_candidate()


@pytest.fixture
def auth_client(candidate):
    client = APIClient()
    client.force_authenticate(user=candidate.user)
    return client


@pytest.fixture
def anon_client():
    return APIClient()


@pytest.fixture(autouse=True)
def reset_throttles():
    """Clear DRF's rate-limit history between tests.

    ``ScopedRateThrottle`` keeps its per-user history in the default cache, which is
    process-wide LocMemCache and is *not* rolled back with the database. Without
    this, the agent-triggering endpoints spend their shared 20/hour budget across
    the whole session and whichever tests happen to run last get a 429 instead of
    the status they assert on - a failure that moves when tests are reordered.
    """
    from django.core.cache import cache

    cache.clear()
    yield
    cache.clear()


@pytest.fixture(autouse=True)
def media_to_tmp(settings, tmp_path):
    """Keep uploaded test files out of backend/media/."""
    settings.MEDIA_ROOT = tmp_path / "media"
    return settings.MEDIA_ROOT


# --- agent / pipeline fixtures ------------------------------------------------

# What a well-behaved analyzer run returns. Kept here so the task tests, the API
# tests and the serializer tests all assert against one shape.
MATCH_RESULT = {
    "match_score": 78,
    "reasoning": "You match the Django and Celery requirements directly. Kubernetes is absent.",
    "matched_skills": ["Python", "Django", "Celery", "PostgreSQL"],
    "missing_skills": ["Kubernetes", "Terraform"],
}


@pytest.fixture(autouse=True)
def no_llm_network(monkeypatch):
    """Hard-fail any test that would really call an LLM.

    An unmocked agent call would otherwise pass silently on a machine with Ollama
    running and hang for three minutes on one without.

    Both backends have to be blocked, and not in the same way: Ollama goes through
    ``requests``, while the Gemini SDK uses httpx internally and never touches it.

    The Gemini side is blocked at the SDK's own method rather than at
    ``gemini_client._generate``, so that the client's error-translation tests can
    still drive ``_generate`` for real against a mocked SDK. Blocking our own seam
    would make the code that turns a dead socket into an ``AgentError`` untestable.
    """

    def blocked(*args, **kwargs):
        raise AssertionError("A test tried to make a real HTTP call. Mock the agent layer.")

    monkeypatch.setattr(requests, "post", blocked)
    monkeypatch.setattr("google.genai.models.Models.generate_content", blocked)


@pytest.fixture
def resume(candidate):
    return Resume.objects.create(
        candidate=candidate,
        file="resumes/2026/07/jane.pdf",
        parsed_text="\n".join(SAMPLE_RESUME_LINES),
    )


@pytest.fixture
def job_description(candidate):
    return JobDescription.objects.create(
        candidate=candidate,
        title="Senior Python Engineer",
        company="Globex",
        raw_text=SAMPLE_JD_TEXT,
    )


@pytest.fixture
def analysis(candidate, resume, job_description):
    """A freshly created, still-pending analysis row."""
    return MatchAnalysis.objects.create(
        candidate=candidate, resume=resume, job_description=job_description
    )


@pytest.fixture
def stub_analyzer(mocker):
    """Replace the LLM call the pipeline task makes. Returns the mock."""
    return mocker.patch("resumes.tasks.analyze_match", return_value=dict(MATCH_RESULT))


@pytest.fixture
def eager_tasks(monkeypatch):
    """Run dispatched tasks inline so an endpoint test sees the finished row.

    Patches each view module's ``run_task`` reference rather than settings: the
    thread runner would otherwise race the assertions, and Celery's eager mode still
    wants a broker configured.
    """
    from agents.tasks import run_task_eager

    monkeypatch.setattr("resumes.views.run_task", run_task_eager)
    monkeypatch.setattr("interviews.views.run_task", run_task_eager)


# --- interview / evaluation fixtures ------------------------------------------

# What a well-behaved question_generator run returns: the four categories, cycled,
# as agents/question_generator.py prescribes.
GENERATED_QUESTIONS = [
    {
        "text": "You migrated a monolith to Celery at Acme - what broke first?",
        "category": "technical",
        "focus": "Celery migration pitfalls",
    },
    {
        "text": "Your resume says 2M requests per day. Walk me through where the bottleneck was.",
        "category": "experience",
        "focus": "scaling a Django API",
    },
    {
        "text": "Tell me about a time you disagreed with a senior engineer's design.",
        "category": "behavioural",
        "focus": "handling technical disagreement",
    },
    {
        "text": "This role needs Kubernetes, which your resume does not show. How would you pick it up?",
        "category": "gap",
        "focus": "Kubernetes readiness",
    },
    {
        "text": "How would you tune a PostgreSQL query that has started doing a sequential scan?",
        "category": "technical",
        "focus": "PostgreSQL query tuning",
    },
]

# What a well-behaved evaluator run returns for one answer.
ANSWER_EVALUATION = {
    "score": 72,
    "verdict": "You named the specific failure and what you changed, which is what makes this land.",
    "strengths": ["Named the real service", "Gave a concrete number"],
    "improvements": ["Say how long the migration took", "Name who else was involved"],
    "model_answer": "The first thing that broke was task idempotency. I would open with that.",
}

# What a well-behaved report run returns over a whole session.
SESSION_REPORT = {
    "overall_score": 68,
    "headline": "Strong on specifics, vague on collaboration.",
    "summary": "You answer technical questions with real numbers. The behavioural answers stay abstract.",
    "strengths": ["Concrete technical detail", "Clear ownership of past work"],
    "priorities": ["Prepare two collaboration stories", "Name the Kubernetes gap head-on"],
    "readiness": "nearly ready",
}


@pytest.fixture
def session(candidate, resume, job_description):
    """A freshly created, still-pending interview session."""
    from interviews.models import InterviewSession

    return InterviewSession.objects.create(
        candidate=candidate, resume=resume, job_description=job_description
    )


@pytest.fixture
def open_session(session):
    """A session whose questions have been written - ready to be answered."""
    session.mark_complete(GENERATED_QUESTIONS)
    return session


@pytest.fixture
def stub_question_generator(mocker):
    """Replace the LLM call the question task makes. Returns the mock."""
    return mocker.patch(
        "interviews.tasks.generate_questions", return_value=[dict(q) for q in GENERATED_QUESTIONS]
    )


@pytest.fixture
def stub_answer_evaluator(mocker):
    """Replace the Gemini call the per-answer task makes. Returns the mock."""
    return mocker.patch(
        "evaluations.tasks.evaluate_answer", return_value=dict(ANSWER_EVALUATION)
    )


@pytest.fixture
def stub_report_builder(mocker):
    """Replace the Gemini call the report task makes. Returns the mock."""
    return mocker.patch("evaluations.tasks.build_report", return_value=dict(SESSION_REPORT))


@pytest.fixture
def answered_session(open_session):
    """A session with every question answered and every answer scored.

    Built through the models rather than the endpoints so report tests do not depend
    on the answer endpoint working.
    """
    from evaluations.models import AnswerEvaluation
    from interviews.models import Answer

    for question in open_session.questions.all():
        answer = Answer.objects.create(
            question=question, text=f"My answer to question {question.order}. " * 4
        )
        AnswerEvaluation.objects.create(answer=answer).mark_complete(dict(ANSWER_EVALUATION))

    return open_session
